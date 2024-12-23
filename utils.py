import pandas as pd
import os
import shutil
from docx import Document
from docx.shared import Inches
import openai
import markdown
from bs4 import BeautifulSoup
import requests



# read excel file and generate docx object
def read_excel_and_generate_word(input_file):

    # 读取 Excel 文件
    df = pd.read_excel(input_file)
    query_column = '装备名称'  # 装备名称列

    if query_column not in df.columns:
        raise ValueError(f"Excel 文件中未找到列 '{query_column}'，请检查列名是否正确！")
    
    # 提取装备名称
    equipment_names = df[query_column]

    # 创建 Word 文档
    doc = Document()
    doc.add_heading('装备数据报告', level=0)

    return equipment_names, doc


# chat with your model(openai)
def openai_chat(model, prompt):
    return openai.ChatCompletion.create(
                model=model,
                messages=[
                    {"role": "system", "content": "你是一个军事领域专家。请回答用户给你的装备相关问题，请你输出纯文本（txt），而不含任何markdown格式。"},
                    {"role": "user", "content": prompt+"\n 注意 不要以markdown格式输出。"}
                ]
            )["choices"][0]["message"]["content"]

# if the result still have markdown formatting, remove it. you can use the controller to control whether to remove markdown formatting.
# if controller is True, remove markdown formatting.
def remove_markdown_formatting(md, controller:bool):
    if not controller:
        return md

    html = markdown.markdown(md)
    soup = BeautifulSoup(html, "html.parser")

    return soup.get_text()

# write word
def write_word(doc, title, context_jane, context_wiki):
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"数据源 1（简氏防务）：{remove_markdown_formatting(context_jane, True)}")
    doc.add_paragraph(f"数据源 2（维基百科）：{remove_markdown_formatting(context_wiki, True)}")

# get wiki first picture
def get_wiki_first_pic(page):
    for image in page.images:
        if image.endswith(".jpg"):
            return image
        

def save_image_from_url(url, save_path):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }  # 伪装为常见的浏览器
    try:
        response = requests.get(url, headers=headers, stream=True)  # 添加 User-Agent
        if response.status_code == 200:
            with open(save_path, "wb") as file:
                for chunk in response.iter_content(1024):  # 分块写入
                    file.write(chunk)
            print(f"图片已成功保存到: {save_path}")
        else:
            print(f"请求失败，状态码: {response.status_code}")
    except Exception as e:
        print(f"发生错误: {e}")


def write_pic_to_word(doc, pic_url):

    doc.add_picture(pic_url, width=Inches(6))